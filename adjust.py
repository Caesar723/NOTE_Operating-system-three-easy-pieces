from docx import Document
import os
import re

Topics=("Virtualisation","Concurrency","Persistence","Distribution")


class fileHandle:

    def __init__(self,topic:str) -> None:
        self.path="/Users/chenxuanpei/Desktop/code/Operation_System/"+topic

    def Path(self,name):
        return f'{self.path}/{name}'

def checkTitle(text,comp=re.compile(r'\d*[\.。,]\d* *$'))->bool:
    return bool(comp.match(text))

    
def sortElement(element):
    try:
        return int(element)
    except:
        return 999999999

def addTxt(handle,name,document):
    first=True
    document.add_heading(name, level=2)
    # p=document.add_paragraph(name)
    # p.style=document.styles['Heading']
    with open(f'{handle.Path(name)}/recode.txt','r') as f:
        getTxt=f.read()
    for each in getTxt.split('\n'):
        if checkTitle(each):
            if first:
                first=False
            else:
                document.add_page_break()
            document.add_heading(each,level=3)
        else:
            document.add_paragraph(each)
    
    document.add_page_break()


def main():
    document = Document()
    for topic in Topics:
        document.add_heading(topic, level=1)
        handle=fileHandle(topic)
        getList=os.listdir(handle.path)
        getList.sort(key=sortElement)
        print(getList)
        for dir in getList[:-1]:

            addTxt(handle,dir,document)

    document.save(r"/Users/chenxuanpei/Desktop/code/Operation_System/Operation_System.docx")

if __name__=="__main__":
    main()
    